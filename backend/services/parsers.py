from __future__ import annotations

import csv
import json
import logging
import re
import subprocess
import uuid
from dataclasses import dataclass
from math import sqrt
from datetime import datetime, timedelta, timezone
from pathlib import Path
from time import sleep

import requests

from backend.core.config import settings
from backend.domain.models import IntermediateDocument, ParagraphSpan, SectionNode
from backend.services.blob_storage import build_blob_artifact_store
from backend.services.foundry_openai import describe_image_with_foundry

logger = logging.getLogger(__name__)

HEADING_ROLE_NAMES = {"title", "sectionheading"}
NON_BODY_ROLE_NAMES = {"pageheader", "pagefooter", "pagenumber"}
PART_HEADING_PATTERN = re.compile(r"^(part|chapter)\b", re.IGNORECASE)
HEADING_NUMBER_PATTERN = re.compile(r"^\d{1,4}\s+")
TABLE_OF_CONTENTS_HEADING_PATTERN = re.compile(r"^(contents|table of contents)\b", re.IGNORECASE)

try:
    from pypdf import PdfReader, PdfWriter
except Exception:  # pragma: no cover - optional dependency
    PdfReader = None  # type: ignore[assignment]
    PdfWriter = None  # type: ignore[assignment]

try:
    import docx  # type: ignore[import-not-found]
except Exception:  # pragma: no cover - optional dependency
    docx = None  # type: ignore[assignment]

try:
    from PIL import Image as PILImage, ImageStat
except Exception:  # pragma: no cover - optional dependency
    PILImage = None  # type: ignore[assignment]
    ImageStat = None  # type: ignore[assignment]

try:
    import pypdfium2 as pdfium  # type: ignore[import-not-found]
except Exception:  # pragma: no cover - optional dependency
    pdfium = None  # type: ignore[assignment]

# Figure rendering parameters. Engineering PDFs frequently draw figures as 1-bit
# image masks / soft-masked (SMask) stencils whose paint colour lives in the page
# content stream, not in the embedded XObject. Pulling the raw XObject therefore
# yields a solid-black bitmap, so we render the composited page region instead.
_FIGURE_RENDER_DPI = 150.0
_FIGURE_CLUSTER_GAP_PTS = 6.0
_FIGURE_MIN_DIMENSION_PTS = 24.0
_FIGURE_REGION_PADDING_PTS = 4.0

# Vector-figure fallback. Business and analyst decks (e.g. slide-style reports)
# draw charts as native vector graphics - hundreds of path/form objects with no
# embedded raster image to anchor on. When a page has no raster figure region but
# substantial vector drawing content, we cluster the vector object bounds into
# chart regions and render those instead, so the chart still becomes a thumbnail.
_FIGURE_VECTOR_OBJECT_THRESHOLD = 8
_FIGURE_VECTOR_MIN_AREA_RATIO = 0.05
_FIGURE_VECTOR_MIN_DIMENSION_PTS = 72.0
_FIGURE_VECTOR_MAX_REGIONS_PER_PAGE = 2
# A vector region that spans almost the entire page height or width is a template
# band (text rail, page border, background panel), not a chart, so it is rejected.
_FIGURE_VECTOR_MAX_EXTENT_RATIO = 0.95
# pdfium page object type codes for vector drawing primitives (path, shading, form).
_FIGURE_VECTOR_OBJECT_TYPES = (2, 4, 5)


@dataclass(slots=True)
class DocumentProfile:
    format: str
    complexity: str
    page_count: int | None
    parser_path: str
    warnings: list[str]


@dataclass(slots=True)
class PdfSegment:
    path: Path
    page_start: int
    page_end: int


@dataclass(slots=True)
class ParagraphBlock:
    text: str
    page_start: int | None = None
    page_end: int | None = None
    role: str | None = None


def _resampling_filter() -> object:
    if PILImage is None:
        raise RuntimeError("Pillow is required for image normalization.")
    resampling = getattr(PILImage, "Resampling", PILImage)
    return resampling.LANCZOS


def _save_normalized_pil(
    pil_image: object,
    artifact_base: Path,
    *,
    base_metadata: dict[str, object] | None = None,
) -> tuple[Path, dict[str, object]]:
    metadata: dict[str, object] = dict(base_metadata or {})
    normalized = pil_image
    width, height = normalized.size
    metadata["original_width"] = width
    metadata["original_height"] = height
    pixel_count = width * height
    metadata["original_pixel_count"] = pixel_count

    scale_candidates = [1.0]
    if settings.max_figure_image_pixels > 0 and pixel_count > settings.max_figure_image_pixels:
        scale_candidates.append(sqrt(settings.max_figure_image_pixels / pixel_count))
    if settings.max_figure_image_dimension > 0:
        largest_dimension = max(width, height)
        if largest_dimension > settings.max_figure_image_dimension:
            scale_candidates.append(settings.max_figure_image_dimension / largest_dimension)
    scale = min(scale_candidates)
    if scale < 1.0:
        resized_width = max(1, int(width * scale))
        resized_height = max(1, int(height * scale))
        normalized = normalized.resize((resized_width, resized_height), _resampling_filter())
        metadata["resized_width"] = resized_width
        metadata["resized_height"] = resized_height
    else:
        metadata["resized_width"] = width
        metadata["resized_height"] = height

    if normalized.mode not in {"RGB", "RGBA"}:
        normalized = normalized.convert("RGBA" if "A" in normalized.getbands() else "RGB")
    output_format = "PNG"
    artifact_path = artifact_base.with_suffix(".png")
    normalized.save(artifact_path, format=output_format, optimize=True)
    metadata["normalized_image"] = True
    metadata["output_format"] = output_format.lower()
    metadata["artifact_pixel_count"] = metadata["resized_width"] * metadata["resized_height"]
    metadata["downscaled"] = bool(scale < 1.0)
    return artifact_path, metadata


def _normalize_figure_image(
    image: object,
    artifact_base: Path,
) -> tuple[Path, dict[str, object]]:
    original_name = str(getattr(image, "name", artifact_base.name))
    original_suffix = Path(original_name).suffix.lower() or ".bin"
    metadata: dict[str, object] = {"original_image_name": original_name}

    pil_image = getattr(image, "image", None)
    if PILImage is None or pil_image is None:
        artifact_path = artifact_base.with_suffix(original_suffix)
        artifact_path.write_bytes(getattr(image, "data", b""))
        metadata["normalized_image"] = False
        metadata["output_format"] = original_suffix.lstrip(".")
        return artifact_path, metadata

    return _save_normalized_pil(pil_image.copy(), artifact_base, base_metadata=metadata)


def _cluster_boxes(
    boxes: list[tuple[float, float, float, float]],
    gap: float,
) -> list[tuple[float, float, float, float]]:
    """Merge image-object bounding boxes that sit within ``gap`` points of each
    other into a single figure region, so a chart composed of many stencils is
    captured as one thumbnail rather than dozens of fragments."""
    count = len(boxes)
    parent = list(range(count))

    def find(index: int) -> int:
        while parent[index] != index:
            parent[index] = parent[parent[index]]
            index = parent[index]
        return index

    def union(left: int, right: int) -> None:
        parent[find(left)] = find(right)

    def near(a: tuple[float, float, float, float], b: tuple[float, float, float, float]) -> bool:
        al, ab, ar, at = a
        bl, bb, br, bt = b
        return not (ar + gap < bl or br + gap < al or at + gap < bb or bt + gap < ab)

    for i in range(count):
        for j in range(i + 1, count):
            if near(boxes[i], boxes[j]):
                union(i, j)

    groups: dict[int, list[float]] = {}
    for i in range(count):
        root = find(i)
        left, bottom, right, top = boxes[i]
        if root not in groups:
            groups[root] = [left, bottom, right, top]
        else:
            existing = groups[root]
            existing[0] = min(existing[0], left)
            existing[1] = min(existing[1], bottom)
            existing[2] = max(existing[2], right)
            existing[3] = max(existing[3], top)
    return [tuple(values) for values in groups.values()]  # type: ignore[misc]


def _is_blank_image(image: object) -> bool:
    """True when a rendered crop is essentially uniform white or black, which
    indicates whitespace or a failed region rather than a real figure."""
    if ImageStat is None:
        return False
    try:
        stat = ImageStat.Stat(image.convert("RGB"))
        mean = sum(stat.mean) / 3.0
        spread = max(stat.stddev) if stat.stddev else 0.0
    except Exception:
        return False
    if mean >= 252.0 and spread < 3.0:
        return True
    if mean <= 3.0 and spread < 3.0:
        return True
    return False


def _finalize_figure_artifact(
    figure: dict[str, object],
    artifact_path: Path,
    *,
    doc_id: str,
    source_name: str,
    page_number: int,
    blob_store: object,
    describe_images: bool,
) -> None:
    if blob_store is not None:
        blob_name = f"documents/{doc_id}/figures/{artifact_path.name}"
        try:
            upload_result = blob_store.upload_file(
                artifact_path,
                blob_name=blob_name,
                metadata={
                    "docid": doc_id,
                    "source": source_name[:120],
                    "page": str(page_number),
                },
            )
            figure.update(upload_result)
        except Exception as exc:
            logger.warning(
                "blob upload failed for figure artifact",
                extra={"context": {"artifact_path": str(artifact_path), "error": str(exc)}},
            )
            figure["blob_upload_error"] = str(exc)
    if describe_images:
        try:
            description, model_endpoint = describe_image_with_foundry(
                artifact_path,
                source_name=source_name,
                page_number=page_number,
            )
            figure["description"] = description
            figure["description_model_endpoint"] = model_endpoint
        except Exception as exc:
            error_text = str(exc)
            if "content_filter" in error_text:
                reason = "content_filter_jailbreak" if "jailbreak" in error_text else "content_filter"
            else:
                reason = "error"
            logger.warning(
                "image description failed",
                extra={
                    "context": {
                        "artifact_path": str(artifact_path),
                        "page_number": page_number,
                        "reason": reason,
                        "error": error_text[:600],
                    }
                },
            )
            figure["description_error"] = error_text


def _extract_pdf_figures_rendered(
    path: Path,
    doc_id: str,
    source_name: str,
    *,
    figure_dir: Path,
    blob_store: object,
    describe_images: bool,
    max_artifacts: int,
) -> list[dict[str, object]]:
    """Render each PDF page and crop figure regions from the composited bitmap.

    Unlike raw XObject extraction, this faithfully reproduces image masks,
    soft-masked (SMask) transparency, CMYK/indexed colour spaces, and vector
    overlays, so engineering figures no longer come out solid black."""
    figures: list[dict[str, object]] = []
    document = pdfium.PdfDocument(str(path))
    scale = _FIGURE_RENDER_DPI / 72.0
    try:
        for page_index in range(len(document)):
            page = document[page_index]
            page_number = page_index + 1
            try:
                page_width, page_height = page.get_size()
            except Exception:
                continue
            try:
                page_objects = list(page.get_objects(max_depth=4))
            except Exception as exc:
                logger.warning(
                    "pdfium page object enumeration failed",
                    extra={"context": {"source": source_name, "page_number": page_number, "error": str(exc)}},
                )
                continue
            image_boxes: list[tuple[float, float, float, float]] = []
            vector_boxes: list[tuple[float, float, float, float]] = []
            for obj in page_objects:
                try:
                    object_type = obj.type
                except Exception:
                    continue
                if object_type == pdfium.raw.FPDF_PAGEOBJ_IMAGE:
                    bucket = image_boxes
                elif object_type in _FIGURE_VECTOR_OBJECT_TYPES:
                    bucket = vector_boxes
                else:
                    continue
                try:
                    left, bottom, right, top = obj.get_bounds()
                except Exception:
                    continue
                if right - left <= 0 or top - bottom <= 0:
                    continue
                bucket.append((left, bottom, right, top))

            # Primary path: regions anchored on embedded raster images (the common
            # case for engineering and scanned documents).
            regions: list[tuple[tuple[float, float, float, float], str]] = []
            if image_boxes:
                for box in _cluster_boxes(image_boxes, _FIGURE_CLUSTER_GAP_PTS):
                    box_left, box_bottom, box_right, box_top = box
                    if (box_right - box_left) >= _FIGURE_MIN_DIMENSION_PTS and (
                        box_top - box_bottom
                    ) >= _FIGURE_MIN_DIMENSION_PTS:
                        regions.append((box, "page_render"))

            # Fallback path: vector-drawn charts (analyst/slide decks) carry no raster
            # image to anchor on, so cluster the vector primitives into chart regions.
            if not regions and len(vector_boxes) >= _FIGURE_VECTOR_OBJECT_THRESHOLD:
                page_area = max(1.0, page_width * page_height)
                sized_regions: list[tuple[float, tuple[float, float, float, float]]] = []
                for box in _cluster_boxes(vector_boxes, _FIGURE_CLUSTER_GAP_PTS):
                    box_left, box_bottom, box_right, box_top = box
                    region_width = box_right - box_left
                    region_height = box_top - box_bottom
                    if (
                        region_width < _FIGURE_VECTOR_MIN_DIMENSION_PTS
                        or region_height < _FIGURE_VECTOR_MIN_DIMENSION_PTS
                    ):
                        continue
                    if (
                        region_height >= page_height * _FIGURE_VECTOR_MAX_EXTENT_RATIO
                        or region_width >= page_width * _FIGURE_VECTOR_MAX_EXTENT_RATIO
                    ):
                        continue
                    region_area = region_width * region_height
                    if region_area / page_area < _FIGURE_VECTOR_MIN_AREA_RATIO:
                        continue
                    sized_regions.append((region_area, box))
                sized_regions.sort(key=lambda item: item[0], reverse=True)
                for _, box in sized_regions[:_FIGURE_VECTOR_MAX_REGIONS_PER_PAGE]:
                    regions.append((box, "page_render_vector"))

            if not regions:
                continue
            try:
                bitmap = page.render(scale=scale)
                page_image = bitmap.to_pil().convert("RGB")
            except Exception as exc:
                logger.warning(
                    "pdfium page render failed",
                    extra={"context": {"source": source_name, "page_number": page_number, "error": str(exc)}},
                )
                continue
            page_px_width, page_px_height = page_image.size
            group_index = 0
            for (left, bottom, right, top), extraction_method in regions:
                pad = _FIGURE_REGION_PADDING_PTS
                x0 = int(round((left - pad) * scale))
                x1 = int(round((right + pad) * scale))
                y0 = int(round((page_height - (top + pad)) * scale))
                y1 = int(round((page_height - (bottom - pad)) * scale))
                x0, x1 = sorted((max(0, min(page_px_width, x0)), max(0, min(page_px_width, x1))))
                y0, y1 = sorted((max(0, min(page_px_height, y0)), max(0, min(page_px_height, y1))))
                if x1 - x0 < 8 or y1 - y0 < 8:
                    continue
                crop = page_image.crop((x0, y0, x1, y1))
                if _is_blank_image(crop):
                    continue
                group_index += 1
                artifact_base = figure_dir / f"page_{page_number:04d}_figure_{group_index}"
                try:
                    artifact_path, image_metadata = _save_normalized_pil(crop, artifact_base)
                except Exception:
                    continue
                figure: dict[str, object] = {
                    "artifact_id": uuid.uuid4().hex[:12],
                    "page_number": page_number,
                    "image_name": f"page_{page_number:04d}_figure_{group_index}.png",
                    "artifact_path": str(artifact_path),
                    "extraction_method": extraction_method,
                    "region_left": round(left, 2),
                    "region_bottom": round(bottom, 2),
                    "region_right": round(right, 2),
                    "region_top": round(top, 2),
                }
                figure.update(image_metadata)
                _finalize_figure_artifact(
                    figure,
                    artifact_path,
                    doc_id=doc_id,
                    source_name=source_name,
                    page_number=page_number,
                    blob_store=blob_store,
                    describe_images=describe_images,
                )
                figures.append(figure)
                if max_artifacts > 0 and len(figures) >= max_artifacts:
                    logger.info(
                        "parser figure extraction reached document limit",
                        extra={
                            "context": {
                                "doc_id": doc_id,
                                "source_name": source_name,
                                "artifact_limit": max_artifacts,
                            }
                        },
                    )
                    return figures
    finally:
        try:
            document.close()
        except Exception:
            pass
    return figures


def _extract_pdf_figure_artifacts(
    path: Path,
    doc_id: str,
    source_name: str,
    *,
    describe_images: bool | None = None,
    max_artifacts: int | None = None,
) -> list[dict[str, object]]:
    if describe_images is None:
        describe_images = settings.parser_image_understanding_enabled
    if max_artifacts is None:
        max_artifacts = settings.parser_figure_max_artifacts

    figure_dir = settings.artifacts_dir / f"{doc_id}_figures"
    figure_dir.mkdir(parents=True, exist_ok=True)
    blob_store = build_blob_artifact_store()

    if pdfium is not None and PILImage is not None:
        try:
            return _extract_pdf_figures_rendered(
                path,
                doc_id,
                source_name,
                figure_dir=figure_dir,
                blob_store=blob_store,
                describe_images=describe_images,
                max_artifacts=max_artifacts,
            )
        except Exception as exc:  # pragma: no cover - defensive fallback
            logger.warning(
                "pdfium figure rendering failed; using embedded image fallback",
                extra={"context": {"source": source_name, "error": str(exc)}},
            )

    return _extract_pdf_figures_embedded(
        path,
        doc_id,
        source_name,
        figure_dir=figure_dir,
        blob_store=blob_store,
        describe_images=describe_images,
        max_artifacts=max_artifacts,
    )


def _extract_pdf_figures_embedded(
    path: Path,
    doc_id: str,
    source_name: str,
    *,
    figure_dir: Path,
    blob_store: object,
    describe_images: bool,
    max_artifacts: int,
) -> list[dict[str, object]]:
    if PdfReader is None:
        return []
    try:
        reader = PdfReader(str(path))
    except Exception:
        return []

    figures: list[dict[str, object]] = []
    artifact_limit_reached = False
    for page_number, page in enumerate(reader.pages, start=1):
        page_images = getattr(page, "images", None)
        if page_images is None:
            continue
        try:
            image_ids = list(page_images.keys())
        except Exception as exc:
            logger.warning(
                "page image enumeration failed",
                extra={"context": {"source": source_name, "page_number": page_number, "error": str(exc)}},
            )
            continue
        for image_index, image_id in enumerate(image_ids, start=1):
            try:
                image = page_images[image_id]
            except Exception as exc:
                logger.warning(
                    "skipping PDF figure image",
                    extra={
                        "context": {
                            "source": source_name,
                            "page_number": page_number,
                            "image_index": image_index,
                            "image_id": image_id if isinstance(image_id, str) else str(image_id),
                            "error": str(exc),
                        }
                    },
                )
                continue
            artifact_base = figure_dir / f"page_{page_number:04d}_figure_{image_index}"
            try:
                artifact_path, image_metadata = _normalize_figure_image(image, artifact_base)
            except Exception:
                continue
            artifact_id = uuid.uuid4().hex[:12]
            figure: dict[str, object] = {
                "artifact_id": artifact_id,
                "page_number": page_number,
                "image_name": getattr(image, "name", f"figure_{image_index}{artifact_path.suffix}"),
                "artifact_path": str(artifact_path),
            }
            figure.update(image_metadata)
            if blob_store is not None:
                blob_name = f"documents/{doc_id}/figures/{artifact_path.name}"
                try:
                    upload_result = blob_store.upload_file(
                        artifact_path,
                        blob_name=blob_name,
                        metadata={
                            "docid": doc_id,
                            "source": source_name[:120],
                            "page": str(page_number),
                        },
                    )
                    figure.update(upload_result)
                except Exception as exc:
                    logger.warning(
                        "blob upload failed for figure artifact",
                        extra={"context": {"artifact_path": str(artifact_path), "error": str(exc)}},
                    )
                    figure["blob_upload_error"] = str(exc)
            if describe_images:
                try:
                    description, model_endpoint = describe_image_with_foundry(
                        artifact_path,
                        source_name=source_name,
                        page_number=page_number,
                    )
                    figure["description"] = description
                    figure["description_model_endpoint"] = model_endpoint
                except Exception as exc:
                    error_text = str(exc)
                    if "content_filter" in error_text:
                        reason = "content_filter_jailbreak" if "jailbreak" in error_text else "content_filter"
                    else:
                        reason = "error"
                    logger.warning(
                        "image description failed",
                        extra={
                            "context": {
                                "artifact_path": str(artifact_path),
                                "page_number": page_number,
                                "reason": reason,
                                "error": error_text[:600],
                            }
                        },
                    )
                    figure["description_error"] = error_text
            figures.append(figure)
            if max_artifacts > 0 and len(figures) >= max_artifacts:
                artifact_limit_reached = True
                logger.info(
                    "parser figure extraction reached document limit",
                    extra={
                        "context": {
                            "doc_id": doc_id,
                            "source_name": source_name,
                            "artifact_limit": max_artifacts,
                        }
                    },
                )
                break
        if artifact_limit_reached:
            break
    return figures


class ParserSelection:
    def profile(self, path: Path) -> DocumentProfile:
        suffix = path.suffix.lower()
        warnings: list[str] = []
        if suffix in {".txt", ".md", ".json", ".csv"}:
            return DocumentProfile(suffix.lstrip("."), "simple", 1, "local_simple_parser", warnings)
        if suffix in {".pdf"}:
            page_count = self._pdf_page_count(path)
            file_size_bytes = self._file_size_bytes(path)
            complexity = "complex" if (page_count or 0) > 10 else "moderate"
            split_reasons: list[str] = []
            if page_count and page_count > settings.hard_page_split_threshold:
                split_reasons.append(f"{settings.hard_page_split_threshold} pages")
            if file_size_bytes > settings.hard_file_split_threshold_bytes:
                split_reasons.append(f"{settings.hard_file_split_threshold_mb} MB")
            if split_reasons:
                warnings.append(
                    "PDF exceeds the segmentation threshold "
                    + " and ".join(split_reasons)
                    + " and will be segmented before analysis."
                )
            if settings.azure_docint_enabled:
                parser_path = "azure_document_intelligence"
            elif settings.workshop_strict_mode:
                parser_path = "strict_configuration_error"
                warnings.append(
                    "Workshop strict mode requires Azure Document Intelligence for PDF ingestion. "
                    "Configure AZURE_DOCUMENT_INTELLIGENCE_ENDPOINT and AZURE_DOCUMENT_INTELLIGENCE_KEY."
                )
            else:
                parser_path = "fallback_pdf_stub"
            return DocumentProfile("pdf", complexity, page_count, parser_path, warnings)
        if suffix in {".docx", ".pptx", ".xlsx"}:
            if settings.azure_content_understanding_enabled:
                parser_path = "azure_content_understanding"
            elif settings.azure_docint_enabled:
                parser_path = "azure_document_intelligence"
            elif settings.workshop_strict_mode:
                parser_path = "strict_configuration_error"
                warnings.append(
                    "Workshop strict mode requires Azure Content Understanding or Azure Document Intelligence "
                    "for Office documents. Configure Azure parsing before running the workshop."
                )
            else:
                parser_path = "local_office_stub"
            return DocumentProfile(suffix.lstrip("."), "moderate", None, parser_path, warnings)
        if settings.workshop_strict_mode:
            return DocumentProfile(
                "binary",
                "unknown",
                None,
                "strict_configuration_error",
                [
                    "Workshop strict mode supports text, markdown, CSV, JSON, PDF, DOCX, PPTX, and XLSX inputs. "
                    "Convert the file to a supported format or disable WORKSHOP_STRICT_MODE for exploratory use."
                ],
            )
        return DocumentProfile("binary", "unknown", None, "unsupported_fallback", ["Format not fully supported"])

    def _pdf_page_count(self, path: Path) -> int | None:
        if PdfReader is None:
            return None
        try:
            reader = PdfReader(str(path))
            return len(reader.pages)
        except Exception:
            return None

    def _file_size_bytes(self, path: Path) -> int:
        try:
            return path.stat().st_size
        except Exception:
            return 0


class BaseParser:
    parser_path = "base"

    def parse(self, path: Path, doc_id: str, profile: DocumentProfile) -> IntermediateDocument:
        raise NotImplementedError


class AzureCognitiveAuthMixin:
    _token_cache: dict[str, tuple[str, datetime]] = {}

    def _build_cognitive_headers(
        self,
        api_key: str,
        *,
        content_type: str | None = None,
        prefer_aad: bool = False,
    ) -> dict[str, str]:
        headers: dict[str, str] = {}
        if content_type:
            headers["Content-Type"] = content_type
        if api_key and not prefer_aad:
            headers["Ocp-Apim-Subscription-Key"] = api_key
            return headers
        headers["Authorization"] = f"Bearer {self._get_cognitive_services_token()}"
        return headers

    def _retry_with_aad(self, response: requests.Response, prefer_aad: bool) -> bool:
        return response.status_code in {401, 403} and not prefer_aad

    def _get_cognitive_services_token(self) -> str:
        cached = self._token_cache.get("cognitiveservices")
        now = datetime.now(timezone.utc)
        if cached and cached[1] > now + timedelta(minutes=5):
            return cached[0]

        result = subprocess.run(
            [
                settings.azure_cli_path,
                "account",
                "get-access-token",
                "--resource",
                "https://cognitiveservices.azure.com/",
                "--output",
                "json",
            ],
            check=True,
            capture_output=True,
            text=True,
        )
        payload = json.loads(result.stdout)
        token = payload["accessToken"]
        expires_on_epoch = payload.get("expires_on")
        expires_on = payload.get("expiresOn")
        if expires_on_epoch:
            expires_at = datetime.fromtimestamp(int(expires_on_epoch), tz=timezone.utc)
        elif expires_on:
            parsed = datetime.fromisoformat(expires_on.replace("Z", "+00:00"))
            expires_at = parsed if parsed.tzinfo else parsed.replace(tzinfo=timezone.utc)
        else:
            expires_at = now + timedelta(minutes=50)
        self._token_cache["cognitiveservices"] = (token, expires_at)
        return token


class LocalSimpleParser(BaseParser):
    parser_path = "local_simple_parser"

    def parse(self, path: Path, doc_id: str, profile: DocumentProfile) -> IntermediateDocument:
        suffix = path.suffix.lower()
        if suffix == ".csv":
            return self._parse_csv(path, doc_id, profile)
        if suffix == ".json":
            raw = json.loads(path.read_text(encoding="utf-8"))
            text = json.dumps(raw, indent=2)
        elif suffix == ".md":
            text = path.read_text(encoding="utf-8")
            return self._parse_markdown(path, doc_id, profile, text)
        else:
            text = path.read_text(encoding="utf-8", errors="ignore")
        return IntermediateDocument(
            doc_id=doc_id,
            source_name=path.name,
            source_path=str(path),
            format=profile.format,
            complexity=profile.complexity,
            parser_path=self.parser_path,
            page_count=profile.page_count,
            sections=[SectionNode(heading="Document", paragraphs=[text])],
            warnings=profile.warnings,
        )

    def _parse_markdown(
        self, path: Path, doc_id: str, profile: DocumentProfile, text: str
    ) -> IntermediateDocument:
        root_sections: list[SectionNode] = []
        current = SectionNode(heading="Introduction", level=1)
        for line in text.splitlines():
            heading_match = re.match(r"^(#{1,6})\s+(.*)$", line.strip())
            if heading_match:
                if current.paragraphs or current.heading != "Introduction":
                    root_sections.append(current)
                current = SectionNode(
                    heading=heading_match.group(2).strip(),
                    level=len(heading_match.group(1)),
                )
                continue
            if line.strip():
                current.paragraphs.append(line.strip())
        if current.paragraphs or current.heading != "Introduction":
            root_sections.append(current)
        return IntermediateDocument(
            doc_id=doc_id,
            source_name=path.name,
            source_path=str(path),
            format=profile.format,
            complexity=profile.complexity,
            parser_path=self.parser_path,
            page_count=profile.page_count,
            sections=root_sections or [SectionNode(heading="Document", paragraphs=[text])],
            warnings=profile.warnings,
        )

    def _parse_csv(self, path: Path, doc_id: str, profile: DocumentProfile) -> IntermediateDocument:
        with path.open("r", encoding="utf-8", newline="") as handle:
            reader = csv.reader(handle)
            rows = [row for row in reader]
        heading = rows[0] if rows else ["Columns"]
        body = rows[1:] if len(rows) > 1 else []
        return IntermediateDocument(
            doc_id=doc_id,
            source_name=path.name,
            source_path=str(path),
            format=profile.format,
            complexity=profile.complexity,
            parser_path=self.parser_path,
            page_count=profile.page_count,
            sections=[
                SectionNode(
                    heading=f"CSV:{' / '.join(heading)}",
                    paragraphs=[],
                    tables=[body],
                )
            ],
            warnings=profile.warnings,
        )


class LocalOfficeParser(BaseParser):
    parser_path = "local_office_stub"

    def parse(self, path: Path, doc_id: str, profile: DocumentProfile) -> IntermediateDocument:
        sections: list[SectionNode] = []
        warnings = list(profile.warnings)
        if path.suffix.lower() == ".docx" and docx is not None:
            doc = docx.Document(str(path))
            paragraphs = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]
            sections.append(SectionNode(heading="Word Document", paragraphs=paragraphs))
        else:
            warnings.append(
                "Local Office parsing is limited. Configure Azure Content Understanding or Document Intelligence."
            )
            sections.append(
                SectionNode(
                    heading="Office Document",
                    paragraphs=[
                        "This document type requires Azure parsing for production extraction fidelity.",
                        f"File: {path.name}",
                    ],
                )
            )
        return IntermediateDocument(
            doc_id=doc_id,
            source_name=path.name,
            source_path=str(path),
            format=profile.format,
            complexity=profile.complexity,
            parser_path=self.parser_path,
            page_count=profile.page_count,
            sections=sections,
            warnings=warnings,
        )


class AzureDocumentIntelligenceParser(AzureCognitiveAuthMixin, BaseParser):
    parser_path = "azure_document_intelligence"

    def parse(self, path: Path, doc_id: str, profile: DocumentProfile) -> IntermediateDocument:
        if not settings.azure_docint_enabled:
            if settings.workshop_strict_mode:
                raise RuntimeError(
                    "Workshop strict mode requires Azure Document Intelligence for this document path. "
                    "Configure AZURE_DOCUMENT_INTELLIGENCE_ENDPOINT and AZURE_DOCUMENT_INTELLIGENCE_KEY."
                )
            if path.suffix.lower() == ".pdf":
                return FallbackPdfParser().parse(path, doc_id, profile)
            return LocalOfficeParser().parse(path, doc_id, profile)

        warnings = list(profile.warnings)
        metadata = {"model_id": settings.azure_document_intelligence_model}
        figure_sections = self._figure_sections(path, doc_id, metadata)
        if path.suffix.lower() == ".pdf" and self._should_split_pdf(path, profile):
            segment_size = self._recommended_segment_size(path, profile)
            segments = self._split_pdf(path, segment_size)
            sections: list[SectionNode] = []
            for segment in segments:
                result = self._analyze_document_result(segment.path)
                paragraphs = self._extract_paragraph_blocks(
                    result,
                    page_offset=segment.page_start - 1,
                )
                segment_sections = self._build_structured_sections(
                    paragraphs or [ParagraphBlock("No content returned.", segment.page_start, segment.page_end)],
                    f"Pages {segment.page_start}-{segment.page_end}",
                )
                sections.extend(segment_sections)
            metadata["segment_count"] = len(segments)
            metadata["segment_page_ranges"] = [
                {"page_start": segment.page_start, "page_end": segment.page_end} for segment in segments
            ]
            metadata["segmentation_strategy"] = "pdf_page_segmentation"
            metadata["segmentation_triggers"] = self._segmentation_triggers(path, profile)
            metadata["segment_size_pages"] = segment_size
            warnings.append(
                f"PDF was split into {len(segments)} segments of up to {segment_size} pages before analysis."
            )
        else:
            sections = self._analyze_document(path, profile)
        sections.extend(figure_sections)
        return IntermediateDocument(
            doc_id=doc_id,
            source_name=path.name,
            source_path=str(path),
            format=profile.format,
            complexity=profile.complexity,
            parser_path=self.parser_path,
            page_count=profile.page_count,
            sections=sections or [SectionNode(heading="Document", paragraphs=["No content returned."])],
            warnings=warnings,
            metadata=metadata,
        )

    def _file_size_bytes(self, path: Path) -> int:
        try:
            return path.stat().st_size
        except Exception:
            return 0

    def _segmentation_triggers(self, path: Path, profile: DocumentProfile) -> list[str]:
        triggers: list[str] = []
        if profile.page_count and profile.page_count > settings.hard_page_split_threshold:
            triggers.append("page_count")
        if self._file_size_bytes(path) > settings.hard_file_split_threshold_bytes:
            triggers.append("file_size")
        return triggers

    def _should_split_pdf(self, path: Path, profile: DocumentProfile) -> bool:
        return bool(self._segmentation_triggers(path, profile))

    def _recommended_segment_size(self, path: Path, profile: DocumentProfile) -> int:
        page_limit_segment_size = settings.max_pages_per_segment
        file_size_bytes = self._file_size_bytes(path)
        if not profile.page_count or file_size_bytes <= settings.hard_file_split_threshold_bytes:
            return page_limit_segment_size

        required_segments_by_size = max(
            1,
            -(-file_size_bytes // settings.hard_file_split_threshold_bytes),
        )
        size_based_segment_size = max(1, -(-profile.page_count // required_segments_by_size))
        return min(page_limit_segment_size, size_based_segment_size)

    def _figure_sections(self, path: Path, doc_id: str, metadata: dict[str, object]) -> list[SectionNode]:
        if path.suffix.lower() != ".pdf" or not settings.parser_figure_extraction_enabled:
            return []
        figures = _extract_pdf_figure_artifacts(path, doc_id, path.name)
        if not figures:
            return []
        metadata["figure_count"] = len(figures)
        metadata["figure_artifacts"] = figures
        if settings.parser_figure_max_artifacts > 0 and len(figures) >= settings.parser_figure_max_artifacts:
            metadata["figure_artifacts_truncated"] = True
            metadata["figure_artifact_limit"] = settings.parser_figure_max_artifacts
        paragraphs = [
            (
                f"Figure extracted from page {figure['page_number']}: {figure['image_name']} | "
                f"artifact: {figure['artifact_path']}"
                + (f" | description: {figure['description']}" if figure.get("description") else "")
            )
            for figure in figures
        ]
        return [SectionNode(heading="Extracted Figures", level=1, paragraphs=paragraphs)]

    def _analyze_document(self, path: Path, profile: DocumentProfile) -> list[SectionNode]:
        result = self._analyze_document_result(path)
        paragraphs = self._extract_paragraph_blocks(result)
        return self._build_structured_sections(
            paragraphs or [ParagraphBlock("No content returned.")],
            "Layout Extraction",
        )

    def _analyze_document_result(self, path: Path) -> dict:
        url = (
            f"{settings.azure_document_intelligence_endpoint.rstrip('/')}"
            f"/documentintelligence/documentModels/{settings.azure_document_intelligence_model}:analyze"
            "?api-version=2024-11-30"
        )
        prefer_aad = False
        body = path.read_bytes()
        headers = self._build_cognitive_headers(
            settings.azure_document_intelligence_key,
            content_type="application/octet-stream",
            prefer_aad=prefer_aad,
        )
        response = requests.post(
            url,
            headers=headers,
            data=body,
            timeout=settings.request_timeout_seconds,
        )
        if self._retry_with_aad(response, prefer_aad):
            prefer_aad = True
            headers = self._build_cognitive_headers(
                settings.azure_document_intelligence_key,
                content_type="application/octet-stream",
                prefer_aad=True,
            )
            response = requests.post(
                url,
                headers=headers,
                data=body,
                timeout=settings.request_timeout_seconds,
            )
        response.raise_for_status()
        operation_location = response.headers.get("Operation-Location")
        if operation_location:
            result = self._poll_operation(operation_location, headers)
        else:
            result = response.json()
        return result

    def _extract_paragraphs(self, result: dict) -> list[str]:
        paragraphs: list[str] = []
        for paragraph in result.get("paragraphs", []):
            content = paragraph.get("content", "").strip()
            if content:
                paragraphs.append(content)
        if not paragraphs and result.get("content"):
            paragraphs = [result["content"]]
        return paragraphs

    def _extract_paragraph_blocks(self, result: dict, *, page_offset: int = 0) -> list[ParagraphBlock]:
        blocks: list[ParagraphBlock] = []
        for paragraph in result.get("paragraphs", []):
            content = paragraph.get("content", "").strip()
            if not content:
                continue
            role = self._normalize_paragraph_role(paragraph.get("role"))
            if self._should_drop_paragraph_block(content, role):
                continue
            page_numbers: list[int] = []
            for region in paragraph.get("boundingRegions") or []:
                page_number = region.get("pageNumber")
                if isinstance(page_number, int):
                    page_numbers.append(page_number + page_offset)
            blocks.append(
                ParagraphBlock(
                    text=content,
                    page_start=min(page_numbers) if page_numbers else None,
                    page_end=max(page_numbers) if page_numbers else None,
                    role=role,
                )
            )
        if not blocks and result.get("content"):
            blocks.append(ParagraphBlock(text=str(result["content"]).strip()))
        return blocks

    def _normalize_paragraph_role(self, value: object) -> str | None:
        if not isinstance(value, str) or not value.strip():
            return None
        return re.sub(r"[^a-z]", "", value.lower())

    def _should_drop_paragraph_block(self, text: str, role: str | None) -> bool:
        normalized = " ".join(text.split()).strip()
        if not normalized:
            return True
        if role in NON_BODY_ROLE_NAMES:
            return True
        if re.fullmatch(r"\d{1,4}", normalized):
            return True
        alpha_count = sum(1 for char in normalized if char.isalpha())
        digit_count = sum(1 for char in normalized if char.isdigit())
        if alpha_count == 0 and digit_count <= 8:
            return True
        if alpha_count <= 2 and len(normalized) <= 6 and not normalized.isalpha():
            return True
        if re.fullmatch(r"[\W_]+", normalized):
            return True
        return False

    def _normalize_heading_text(self, text: str) -> str:
        normalized = " ".join(text.replace("\n", " ").split()).strip(" -–—:")
        normalized = HEADING_NUMBER_PATTERN.sub("", normalized)
        return normalized or "Untitled Section"

    def _looks_like_heading_text(self, text: str) -> bool:
        normalized = self._normalize_heading_text(text)
        if not normalized or len(normalized) > 120:
            return False
        if normalized.endswith((".", ";", "?", "!")):
            return False
        if re.search(r"\b\d{1,4}$", normalized):
            return False
        if TABLE_OF_CONTENTS_HEADING_PATTERN.match(normalized):
            return True
        if PART_HEADING_PATTERN.match(normalized):
            return True
        if " ~ " in normalized:
            return False
        words = normalized.split()
        if not 1 <= len(words) <= 12:
            return False
        alpha_count = sum(1 for char in normalized if char.isalpha())
        if alpha_count < max(3, len(normalized) // 3):
            return False
        capitalized_words = sum(1 for word in words if word[:1].isupper())
        uppercase_words = sum(1 for word in words if word.isupper() and len(word) > 1)
        return capitalized_words >= max(1, len(words) - 1) or uppercase_words >= max(1, len(words) // 2)

    def _is_heading_block(self, block: ParagraphBlock) -> bool:
        if block.role in HEADING_ROLE_NAMES:
            return True
        return self._looks_like_heading_text(block.text)

    def _build_structured_sections(
        self,
        paragraphs: list[ParagraphBlock],
        default_heading: str,
    ) -> list[SectionNode]:
        sections: list[SectionNode] = []
        preamble: list[ParagraphBlock] = []
        current: SectionNode | None = None
        for paragraph in paragraphs:
            if self._is_heading_block(paragraph):
                normalized_heading = self._normalize_heading_text(paragraph.text)
                if current is not None and current.heading == normalized_heading and not current.paragraphs:
                    continue
                if current is not None and (current.paragraphs or current.tables or current.children):
                    sections.append(current)
                current = SectionNode(
                    heading=normalized_heading,
                    level=1,
                    paragraphs=[],
                    paragraph_spans=[],
                    page_start=paragraph.page_start,
                    page_end=paragraph.page_end,
                )
                continue
            if current is None:
                preamble.append(paragraph)
            else:
                current.paragraphs.append(paragraph.text)
                current.paragraph_spans.append(
                    ParagraphSpan(page_start=paragraph.page_start, page_end=paragraph.page_end)
                )
                current.page_start = min(
                    [value for value in [current.page_start, paragraph.page_start] if value is not None],
                    default=current.page_start,
                )
                current.page_end = max(
                    [value for value in [current.page_end, paragraph.page_end] if value is not None],
                    default=current.page_end,
                )
        if current is not None:
            sections.append(current)
        if sections:
            if preamble:
                preamble_page_starts = [block.page_start for block in preamble if block.page_start is not None]
                preamble_page_ends = [block.page_end for block in preamble if block.page_end is not None]
                sections.insert(
                    0,
                    SectionNode(
                        heading=default_heading,
                        level=1,
                        paragraphs=[block.text for block in preamble],
                        paragraph_spans=[
                            ParagraphSpan(page_start=block.page_start, page_end=block.page_end)
                            for block in preamble
                        ],
                        page_start=min(preamble_page_starts) if preamble_page_starts else None,
                        page_end=max(preamble_page_ends) if preamble_page_ends else None,
                    ),
                )
            return sections
        page_starts = [block.page_start for block in paragraphs if block.page_start is not None]
        page_ends = [block.page_end for block in paragraphs if block.page_end is not None]
        return [
            SectionNode(
                heading=default_heading,
                paragraphs=[block.text for block in paragraphs],
                paragraph_spans=[
                    ParagraphSpan(page_start=block.page_start, page_end=block.page_end)
                    for block in paragraphs
                ],
                page_start=min(page_starts) if page_starts else None,
                page_end=max(page_ends) if page_ends else None,
            )
        ]

    def _split_pdf(self, path: Path, segment_size: int) -> list[PdfSegment]:
        if PdfReader is None or PdfWriter is None:
            raise RuntimeError("pypdf is required to split oversized PDF files.")

        reader = PdfReader(str(path))
        segment_dir = settings.artifacts_dir / f"{path.stem}_segments"
        segment_dir.mkdir(parents=True, exist_ok=True)

        segments: list[PdfSegment] = []
        total_pages = len(reader.pages)
        for start_index in range(0, total_pages, segment_size):
            end_index = min(start_index + segment_size, total_pages)
            writer = PdfWriter()
            for page_index in range(start_index, end_index):
                writer.add_page(reader.pages[page_index])
            segment_path = segment_dir / f"{path.stem}_pages_{start_index + 1}_{end_index}.pdf"
            with segment_path.open("wb") as handle:
                writer.write(handle)
            segments.append(
                PdfSegment(path=segment_path, page_start=start_index + 1, page_end=end_index)
            )
        return segments

    def _poll_operation(self, operation_location: str, headers: dict[str, str], max_attempts: int = 300) -> dict:
        for _ in range(max_attempts):
            response = requests.get(operation_location, headers=headers, timeout=settings.request_timeout_seconds)
            response.raise_for_status()
            payload = response.json()
            status = payload.get("status", "").lower()
            if status == "succeeded":
                return payload.get("analyzeResult") or payload
            if status == "failed":
                raise RuntimeError(f"Document Intelligence analysis failed: {payload}")
            sleep(1)
        raise TimeoutError("Document Intelligence analysis timed out.")


class AzureContentUnderstandingParser(AzureCognitiveAuthMixin, BaseParser):
    parser_path = "azure_content_understanding"

    def parse(self, path: Path, doc_id: str, profile: DocumentProfile) -> IntermediateDocument:
        if not settings.azure_content_understanding_enabled:
            if settings.workshop_strict_mode:
                raise RuntimeError(
                    "Workshop strict mode requires Azure Content Understanding for this parser path. "
                    "Configure AZURE_CONTENT_UNDERSTANDING_* or select Azure Document Intelligence instead."
                )
            return AzureDocumentIntelligenceParser().parse(path, doc_id, profile)
        url = (
            f"{settings.azure_content_understanding_endpoint.rstrip('/')}"
            f"/contentunderstanding/analyzers/{settings.azure_content_understanding_analyzer_id}:analyze"
            "?api-version=2025-11-01"
        )
        prefer_aad = False
        headers = self._build_cognitive_headers(
            settings.azure_content_understanding_key,
            prefer_aad=prefer_aad,
        )
        body = path.read_bytes()
        files = {"file": (path.name, body)}
        response = requests.post(url, headers=headers, files=files, timeout=settings.request_timeout_seconds)
        if self._retry_with_aad(response, prefer_aad):
            headers = self._build_cognitive_headers(
                settings.azure_content_understanding_key,
                prefer_aad=True,
            )
            files = {"file": (path.name, body)}
            response = requests.post(url, headers=headers, files=files, timeout=settings.request_timeout_seconds)
        response.raise_for_status()
        operation_location = response.headers.get("Operation-Location")
        if operation_location:
            payload = self._poll_operation(operation_location, headers)
        else:
            payload = response.json()
        extracted = payload.get("result", {}).get("contents", [])
        paragraphs = []
        for item in extracted:
            text = item.get("markdown") or item.get("text") or ""
            if text.strip():
                paragraphs.append(text.strip())
        return IntermediateDocument(
            doc_id=doc_id,
            source_name=path.name,
            source_path=str(path),
            format=profile.format,
            complexity=profile.complexity,
            parser_path=self.parser_path,
            page_count=profile.page_count,
            sections=[SectionNode(heading="Content Understanding", paragraphs=paragraphs)],
            warnings=profile.warnings,
            metadata={"analyzer_id": settings.azure_content_understanding_analyzer_id},
        )

    def _poll_operation(self, operation_location: str, headers: dict[str, str]) -> dict:
        for _ in range(60):
            response = requests.get(operation_location, headers=headers, timeout=settings.request_timeout_seconds)
            response.raise_for_status()
            payload = response.json()
            status = payload.get("status", "").lower()
            if status in {"succeeded", "completed"}:
                return payload
            if status in {"failed", "error"}:
                raise RuntimeError(f"Content Understanding analysis failed: {payload}")
            sleep(1)
        raise TimeoutError("Content Understanding analysis timed out.")


class FallbackPdfParser(BaseParser):
    parser_path = "fallback_pdf_stub"

    def parse(self, path: Path, doc_id: str, profile: DocumentProfile) -> IntermediateDocument:
        warnings = list(profile.warnings)
        warnings.append(
            "PDF parsing fallback is limited. Configure Azure Document Intelligence for layout-aware extraction."
        )
        paragraphs = self._extract_pdf_text(path)
        metadata: dict[str, object] = {}
        figure_sections = self._figure_sections(path, doc_id, metadata)
        sections = [
            SectionNode(
                heading="PDF Fallback",
                paragraphs=paragraphs,
            )
        ]
        sections.extend(figure_sections)
        return IntermediateDocument(
            doc_id=doc_id,
            source_name=path.name,
            source_path=str(path),
            format=profile.format,
            complexity=profile.complexity,
            parser_path=self.parser_path,
            page_count=profile.page_count,
            sections=sections,
            metadata=metadata,
            warnings=warnings,
        )

    def _extract_pdf_text(self, path: Path) -> list[str]:
        paragraphs = [
            "The current environment is not configured for Azure parsing.",
            "The ingestion pipeline still produced chunkable placeholder content and metadata.",
        ]
        if PdfReader is None:
            return paragraphs
        try:
            reader = PdfReader(str(path))
            extracted = []
            for page in reader.pages[:10]:
                text = (page.extract_text() or "").strip()
                if text:
                    extracted.append(text)
            return extracted or paragraphs
        except Exception:
            return paragraphs

    def _figure_sections(self, path: Path, doc_id: str, metadata: dict[str, object]) -> list[SectionNode]:
        if not settings.parser_figure_extraction_enabled:
            return []
        figures = _extract_pdf_figure_artifacts(path, doc_id, path.name)
        if not figures:
            return []
        metadata["figure_count"] = len(figures)
        metadata["figure_artifacts"] = figures
        return [
            SectionNode(
                heading="Extracted Figures",
                paragraphs=[
                    f"Figure extracted from page {figure['page_number']}: {figure['image_name']} | artifact: {figure['artifact_path']}"
                    for figure in figures
                ],
            )
        ]


class StrictConfigurationErrorParser(BaseParser):
    parser_path = "strict_configuration_error"

    def parse(self, path: Path, doc_id: str, profile: DocumentProfile) -> IntermediateDocument:
        detail = next((warning for warning in profile.warnings if warning.strip()), "")
        raise RuntimeError(detail or "Workshop strict mode blocked this document because the required parser path is unavailable.")


class ParserRegistry:
    def __init__(self) -> None:
        self.selector = ParserSelection()
        self.parsers: dict[str, BaseParser] = {
            "local_simple_parser": LocalSimpleParser(),
            "local_office_stub": LocalOfficeParser(),
            "azure_document_intelligence": AzureDocumentIntelligenceParser(),
            "azure_content_understanding": AzureContentUnderstandingParser(),
            "fallback_pdf_stub": FallbackPdfParser(),
            "strict_configuration_error": StrictConfigurationErrorParser(),
            "unsupported_fallback": LocalOfficeParser(),
        }

    def detect(self, path: Path) -> DocumentProfile:
        return self.selector.profile(path)

    def parse(self, path: Path, doc_id: str, profile: DocumentProfile) -> IntermediateDocument:
        parser = self.parsers[profile.parser_path]
        logger.info("selected parser", extra={"context": {"parser_path": parser.parser_path, "file": path.name}})
        return parser.parse(path, doc_id, profile)


parser_registry = ParserRegistry()
